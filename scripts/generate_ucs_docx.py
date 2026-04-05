from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None, header_color="1A56DB"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(hdr_cells[i], header_color)

    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = ''
            p = row_cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(row_cells[c_idx], "F0F4FF")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_attr_table(doc, attrs, header_color="2563EB"):
    table = doc.add_table(rows=len(attrs), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(attrs):
        c0 = table.rows[i].cells[0]
        c1 = table.rows[i].cells[1]
        c0.text = ''
        c1.text = ''
        p0 = c0.paragraphs[0]
        run0 = p0.add_run(k)
        run0.bold = True
        run0.font.size = Pt(10)
        set_cell_shading(c0, "EEF2FF")
        p1 = c1.paragraphs[0]
        run1 = p1.add_run(v)
        run1.font.size = Pt(10)
        c0.width = Cm(4.5)
        c1.width = Cm(12)
    return table

def add_flow_table(doc, headers, rows, header_color="059669"):
    return add_styled_table(doc, headers, rows, header_color=header_color)

def add_alt_table(doc, headers, rows, header_color="D97706"):
    return add_styled_table(doc, headers, rows, header_color=header_color)

def add_bullet_list(doc, items, bold_prefix=None):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(10)

# ===== COVER PAGE =====
for _ in range(4):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('USE CASE SPECIFICATION')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Sistem LohParkir')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Sistem Verifikasi dan Manajemen Parkir Berbasis QR Code')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

doc.add_paragraph('')
doc.add_paragraph('')

info_items = [
    ('Versi', '1.0'),
    ('Tanggal', '5 April 2026'),
    ('Referensi', 'SRS LohParkir v1.1 (Revised)'),
]
t = doc.add_table(rows=3, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate(info_items):
    c0 = t.rows[i].cells[0]
    c1 = t.rows[i].cells[1]
    c0.text = ''
    c1.text = ''
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r0 = p0.add_run(f'{k}:')
    r0.bold = True
    r0.font.size = Pt(11)
    p1 = c1.paragraphs[0]
    r1 = p1.add_run(f'  {v}')
    r1.font.size = Pt(11)
    c0.width = Cm(4)
    c1.width = Cm(8)

doc.add_paragraph('')
doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Disusun oleh:')
run.bold = True
run.font.size = Pt(11)

team = [
    '241401055 — Muhammad Rizky Fadhillah',
    '241401007 — Rifki Al Sauqy',
    '241401070 — Yehezkiel Gustav Setiawan Sitanggang',
    '241401076 — Yasmin Assyifa',
    '241401079 — Edric Roland Li',
    '241401136 — Najla Az Zahra Tanjung',
]
for name in team:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    run.font.size = Pt(10)

doc.add_page_break()

# ===== DAFTAR ISI =====
doc.add_heading('Daftar Isi', level=1)
toc_items = [
    '1. Identifikasi Aktor',
    '2. Ringkasan Use Case',
    '3. Use Case Specifications',
    '   UC-01: Scan dan Validasi QR Code',
    '   UC-02: Laporkan Parkir Ilegal',
    '   UC-03: Pembayaran Digital via QRIS',
    '   UC-04: Dashboard Monitoring Real-time',
    '   UC-05: Manajemen Akun Petugas & Badge QR',
    '   UC-06: Riwayat Pembayaran',
    '   UC-07: Manajemen Laporan',
    '   UC-08: Notifikasi Push',
    '4. Matriks Traceability',
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(11)
    if not item.startswith('   '):
        run.bold = True

doc.add_page_break()

# ===== 1. IDENTIFIKASI AKTOR =====
doc.add_heading('1. Identifikasi Aktor', level=1)
doc.add_paragraph('Berikut adalah aktor-aktor yang terlibat dalam sistem LohParkir:')

add_styled_table(doc,
    ['No', 'Aktor', 'Deskripsi', 'Perangkat'],
    [
        ['A1', 'Masyarakat Umum\n(Pengemudi)', 'Pengguna publik yang memindai QR code, melakukan pembayaran parkir, dan melaporkan parkir ilegal. Tidak perlu login.', 'Smartphone\n(Android/iOS)'],
        ['A2', 'Admin Dishub', 'Staff Dinas Perhubungan Kota Medan yang mengelola operasional harian, melihat dashboard, dan menangani laporan.', 'Komputer\n(Web Browser)'],
        ['A3', 'Database Administrator\n(Superadmin)', 'Pengelola database dengan akses penuh CRUD terhadap data resmi (QR code, lokasi, petugas, tarif).', 'Komputer\n(Akses Backend)'],
        ['A4', 'Petugas Parkir\nTerdaftar', 'Petugas parkir resmi yang membawa badge QR dan memiliki akses ke aplikasi minimal dengan tombol panik darurat.', 'Smartphone\n(Android/iOS)'],
        ['A5', 'Sistem', 'Aktor otomatis yang menjalankan fungsi-fungsi seperti validasi, notifikasi, dan pencatatan transaksi.', 'Server Backend'],
    ],
    col_widths=[1, 3, 8, 3]
)

doc.add_paragraph('')

# ===== 2. RINGKASAN USE CASE =====
doc.add_heading('2. Ringkasan Use Case', level=1)

add_styled_table(doc,
    ['ID', 'Nama Use Case', 'Aktor Utama', 'Prioritas', 'REQ Reference'],
    [
        ['UC-01', 'Scan dan Validasi QR Code', 'Masyarakat Umum', 'High', 'REQ-1 s/d REQ-9'],
        ['UC-02', 'Laporkan Parkir Ilegal', 'Masyarakat Umum', 'High', 'REQ-10 s/d REQ-18'],
        ['UC-03', 'Pembayaran Digital via QRIS', 'Masyarakat Umum', 'High', 'REQ-19 s/d REQ-25'],
        ['UC-04', 'Dashboard Monitoring Real-time', 'Admin Dishub', 'High', 'REQ-26 s/d REQ-33'],
        ['UC-05', 'Manajemen Akun Petugas & Badge QR', 'Admin Dishub, Superadmin', 'High', 'REQ-34 s/d REQ-38'],
        ['UC-06', 'Riwayat Pembayaran', 'Masyarakat Umum', 'Low', 'REQ-39 s/d REQ-40'],
        ['UC-07', 'Manajemen Laporan', 'Admin Dishub', 'High', 'REQ-41 s/d REQ-44'],
        ['UC-08', 'Notifikasi Push', 'Sistem', 'Medium', 'REQ-45 s/d REQ-48'],
    ]
)

doc.add_page_break()

# ===== 3. USE CASE SPECIFICATIONS =====
doc.add_heading('3. Use Case Specifications', level=1)

# ----- UC-01 -----
doc.add_heading('UC-01: Scan dan Validasi QR Code', level=2)

add_attr_table(doc, [
    ('Use Case ID', 'UC-01'),
    ('Nama', 'Scan dan Validasi QR Code'),
    ('Deskripsi', 'Masyarakat memindai QR code yang terletak di area parkir atau pada badge petugas parkir. Sistem memvalidasi QR code terhadap database resmi Dishub/Pemda dan menampilkan status keaslian beserta informasi detail.'),
    ('Aktor Utama', 'Masyarakat Umum (A1)'),
    ('Aktor Pendukung', 'Sistem (A5)'),
    ('Prioritas', 'High'),
    ('Trigger', 'Pengguna membuka aplikasi dan menekan tombol "Scan QR".'),
    ('Frekuensi', 'Sangat sering — setiap kali pengguna memarkir kendaraan.'),
])
doc.add_paragraph('')

doc.add_heading('Preconditions', level=3)
add_bullet_list(doc, [
    'Aplikasi LohParkir telah terinstal di smartphone pengguna.',
    'Smartphone memiliki kamera yang berfungsi (minimal 5 MP).',
    'Koneksi internet tersedia (atau mode offline dengan data cache).',
    'Petugas parkir mengenakan badge QR code resmi yang terlihat.',
])

doc.add_heading('Postconditions (Success)', level=3)
add_bullet_list(doc, [
    'Hasil validasi QR code ditampilkan (resmi/tidak resmi).',
    'Informasi petugas ditampilkan: nama, nomor badge, zona, tarif resmi.',
    'Data scan tercatat di database server (tabel scan_history).',
])

doc.add_heading('Postconditions (Failure)', level=3)
add_bullet_list(doc, [
    'Pesan error yang ramah ditampilkan kepada pengguna.',
    'Pengguna diberi opsi untuk mencoba kembali atau melaporkan.',
])

doc.add_heading('Basic Flow (Main Success Scenario)', level=3)
add_flow_table(doc,
    ['Langkah', 'Aktor', 'Sistem'],
    [
        ['1', 'Pengguna membuka aplikasi dan menekan tombol "Scan QR".', '—'],
        ['2', '—', 'Sistem mengaktifkan kamera perangkat.'],
        ['3', 'Pengguna mengarahkan kamera ke QR code pada badge petugas.', '—'],
        ['4', '—', 'Sistem mendeteksi dan mendekode QR code (format: LOHPARKIR-DSH-YYYY-NNN).'],
        ['5', '—', 'Sistem mengirimkan data QR ke server untuk validasi (POST /api/qr/validate).'],
        ['6', '—', 'Server memvalidasi QR code terhadap database resmi.'],
        ['7', '—', 'Sistem menampilkan hasil VALID: nama lokasi parkir (REQ-5), nama petugas (REQ-6), tarif resmi (REQ-7), area operasional (REQ-8), nomor badge dan status aktif/non-aktif (REQ-9).'],
        ['8', 'Pengguna melihat informasi dan memutuskan untuk membayar atau meninggalkan.', '—'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Alternative Flow: AF-01 — QR Code Tidak Valid', level=3)
add_alt_table(doc,
    ['Langkah', 'Deskripsi'],
    [
        ['4a', 'Sistem mendekode QR code namun format sesuai LOHPARKIR-DSH-YYYY-NNN.'],
        ['5a', 'Server tidak menemukan QR code di database.'],
        ['6a', 'Sistem menampilkan peringatan: "QR Code tidak terdaftar di database Dishub" (REQ-10).'],
        ['7a', 'Sistem menampilkan tombol "Laporkan" untuk melaporkan QR palsu (REQ-11). → Berlanjut ke UC-02.'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Alternative Flow: AF-02 — QR Code Format Tidak Sesuai', level=3)
add_alt_table(doc,
    ['Langkah', 'Deskripsi'],
    [
        ['4b', 'QR code terdeteksi namun format tidak sesuai standar LOHPARKIR-DSH-YYYY-NNN.'],
        ['5b', 'Sistem menampilkan pesan: "Format QR code tidak dikenali. Ini bukan QR resmi LohParkir."'],
        ['6b', 'Sistem menampilkan tombol "Laporkan". → Berlanjut ke UC-02.'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Alternative Flow: AF-03 — Petugas Tidak Aktif', level=3)
add_alt_table(doc,
    ['Langkah', 'Deskripsi'],
    [
        ['6c', 'Server menemukan QR code, tetapi status petugas "inactive".'],
        ['7c', 'Sistem menampilkan peringatan: "Petugas ini sudah tidak aktif." dengan detail petugas dan tombol "Laporkan".'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Alternative Flow: AF-04 — Mode Offline', level=3)
add_alt_table(doc,
    ['Langkah', 'Deskripsi'],
    [
        ['5d', 'Pengiriman data ke server gagal karena tidak ada koneksi internet.'],
        ['6d', 'Sistem mengecek cache lokal (AsyncStorage) untuk data offline.'],
        ['7d', 'Jika ditemukan di cache: tampilkan data dari cache dengan label "Data Offline".'],
        ['8d', 'Jika tidak ditemukan di cache: tampilkan pesan "Tidak dapat memvalidasi. Periksa koneksi internet Anda." dengan tombol retry.'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Alternative Flow: AF-05 — Input QR Manual', level=3)
add_alt_table(doc,
    ['Langkah', 'Deskripsi'],
    [
        ['1e', 'Pengguna memilih opsi "Input Manual" alih-alih scan kamera.'],
        ['2e', 'Sistem menampilkan field input teks.'],
        ['3e', 'Pengguna memasukkan kode QR secara manual.'],
        ['4e', 'Lanjut ke langkah 5 dari Basic Flow.'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Exception Flow', level=3)
add_styled_table(doc,
    ['Kode', 'Deskripsi'],
    [
        ['EX-01', 'Kamera tidak tersedia atau izin ditolak → Sistem menampilkan pesan error dan menawarkan input manual.'],
        ['EX-02', 'Server error (5xx) → Sistem menampilkan pesan: "Terjadi gangguan server. Silakan coba lagi nanti."'],
        ['EX-03', 'Timeout (> 2 detik) → Sistem menampilkan pesan timeout dan tombol retry.'],
    ],
    header_color="DC2626"
)
doc.add_paragraph('')

doc.add_heading('Business Rules', level=3)
add_bullet_list(doc, [
    'Format QR code wajib: LOHPARKIR-DSH-YYYY-NNN (BR-03, NFR-07).',
    'Kamera harus mendeteksi QR dalam 1.5 detik (NFR-02).',
    'Respons server maksimal 2 detik (NFR-01).',
    'Pengguna tidak perlu login untuk melakukan scan.',
])

doc.add_page_break()

# ----- UC-02 -----
doc.add_heading('UC-02: Laporkan Parkir Ilegal', level=2)

add_attr_table(doc, [
    ('Use Case ID', 'UC-02'),
    ('Nama', 'Laporkan Parkir Ilegal'),
    ('Deskripsi', 'Masyarakat melaporkan aktivitas parkir ilegal, pungutan liar, atau QR code palsu. Laporan mencakup bukti foto, lokasi GPS otomatis, dan deskripsi pengguna.'),
    ('Aktor Utama', 'Masyarakat Umum (A1)'),
    ('Aktor Pendukung', 'Sistem (A5)'),
    ('Prioritas', 'High'),
    ('Trigger', '(1) Pengguna menekan tombol "Laporkan" setelah scan QR gagal, atau (2) Pengguna membuka tab Laporan dan menekan tombol buat laporan baru.'),
    ('Frekuensi', 'Sering — setiap kali ditemukan parkir ilegal.'),
])
doc.add_paragraph('')

doc.add_heading('Preconditions', level=3)
add_bullet_list(doc, [
    'Aplikasi LohParkir terinstal dan dibuka.',
    'GPS perangkat aktif dan izin lokasi diberikan.',
    'Kamera tersedia untuk mengambil foto bukti.',
    'Koneksi internet tersedia untuk pengiriman laporan.',
])

doc.add_heading('Postconditions (Success)', level=3)
add_bullet_list(doc, [
    'Laporan tersimpan di database server (tabel reports).',
    'Nomor tiket pengaduan di-generate (format: LP-YYYYMMDD-XXXX).',
    'Status awal laporan: "pending".',
    'Notifikasi dikirim ke Admin Dishub (UC-08).',
])

doc.add_heading('Postconditions (Failure)', level=3)
add_bullet_list(doc, [
    'Laporan disimpan secara lokal di perangkat untuk dikirim ulang nanti.',
    'Pesan error ditampilkan dengan instruksi retry.',
])

doc.add_heading('Basic Flow (Main Success Scenario)', level=3)
add_flow_table(doc,
    ['Langkah', 'Aktor', 'Sistem'],
    [
        ['1', 'Pengguna menekan tombol "Laporkan" (dari hasil scan gagal atau dari tab Laporan).', '—'],
        ['2', '—', 'Sistem menampilkan formulir laporan.'],
        ['3', 'Pengguna memilih jenis laporan: "QR Palsu" atau "Parkir Ilegal".', '—'],
        ['4', 'Pengguna menekan tombol "Ambil Foto" dan mengambil foto bukti (REQ-12).', '—'],
        ['5', '—', 'Sistem menerima foto dan menampilkan preview.'],
        ['6', '—', 'Sistem secara otomatis menangkap koordinat GPS (latitude, longitude) (REQ-13).'],
        ['7', 'Pengguna mengisi deskripsi laporan di field teks (REQ-14).', '—'],
        ['8', 'Pengguna menekan tombol "Kirim Laporan".', '—'],
        ['9', '—', 'Sistem mengirim laporan ke server (POST /api/reports).'],
        ['10', '—', 'Server menyimpan laporan dan men-generate nomor tiket (REQ-16).'],
        ['11', '—', 'Sistem menampilkan konfirmasi: "Laporan berhasil dikirim" beserta nomor tiket pengaduan.'],
        ['12', '—', 'Sistem men-trigger notifikasi ke Admin Dishub tentang laporan baru (REQ-46).'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Alternative Flow: AF-01 — Laporan dari Tab Laporan', level=3)
add_alt_table(doc,
    ['Langkah', 'Deskripsi'],
    [
        ['1a', 'Pengguna membuka tab "Laporan" dari navigasi bawah.'],
        ['2a', 'Pengguna menekan tombol "Buat Laporan Baru".'],
        ['3a', 'Lanjut ke langkah 3 dari Basic Flow (tanpa QR code terkait).'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Alternative Flow: AF-02 — Dengan QR Code Terkait', level=3)
add_alt_table(doc,
    ['Langkah', 'Deskripsi'],
    [
        ['1b', 'Pengguna menekan "Laporkan" dari hasil scan QR tidak valid.'],
        ['2b', 'Sistem otomatis mengisi field qr_code_reported dengan kode QR yang di-scan.'],
        ['3b', 'Lanjut ke langkah 3 dari Basic Flow.'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Alternative Flow: AF-03 — Gagal Upload (Tidak Ada Internet)', level=3)
add_alt_table(doc,
    ['Langkah', 'Deskripsi'],
    [
        ['9c', 'Pengiriman laporan ke server gagal.'],
        ['10c', 'Sistem menyimpan laporan secara lokal di perangkat (REQ-18, NFR-17).'],
        ['11c', 'Sistem menampilkan pesan: "Laporan disimpan offline. Akan dikirim otomatis saat koneksi tersedia."'],
        ['12c', 'Saat koneksi pulih, sistem otomatis mengirim laporan yang tertunda.'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Exception Flow', level=3)
add_styled_table(doc,
    ['Kode', 'Deskripsi'],
    [
        ['EX-01', 'GPS tidak aktif → Sistem menampilkan dialog untuk mengaktifkan GPS. Koordinat wajib untuk laporan.'],
        ['EX-02', 'Batas laporan harian tercapai (5 laporan/hari) → Sistem menampilkan pesan: "Anda telah mencapai batas laporan harian." (BR-01).'],
    ],
    header_color="DC2626"
)
doc.add_paragraph('')

doc.add_heading('Business Rules', level=3)
add_bullet_list(doc, [
    'Maksimal 5 laporan per hari per pengguna (BR-01, NFR-11).',
    'Laporan tidak dapat diedit setelah dikirim (BR-08).',
    'Laporan tidak dapat dihapus oleh pengguna (BR-02).',
    'Semua laporan harus dimoderasi oleh admin (NFR-11).',
    'Identitas pelapor harus dilindungi dan anonim (NFR-10).',
])

doc.add_page_break()

# ----- UC-03 -----
doc.add_heading('UC-03: Pembayaran Digital via QRIS', level=2)

add_attr_table(doc, [
    ('Use Case ID', 'UC-03'),
    ('Nama', 'Pembayaran Digital via QRIS'),
    ('Deskripsi', 'Setelah validasi QR berhasil, pengguna dapat membayar tarif parkir resmi secara digital melalui QRIS atau dicatat secara tunai oleh petugas. Semua pembayaran tercatat di database.'),
    ('Aktor Utama', 'Masyarakat Umum (A1)'),
    ('Aktor Pendukung', 'Petugas Parkir (A4), Sistem (A5)'),
    ('Prioritas', 'High'),
    ('Trigger', 'Pengguna menekan tombol "Bayar" setelah hasil validasi QR menunjukkan status resmi.'),
    ('Frekuensi', 'Sangat sering — setiap kali pengguna memarkir kendaraan di lokasi resmi.'),
])
doc.add_paragraph('')

doc.add_heading('Preconditions', level=3)
add_bullet_list(doc, [
    'QR code telah divalidasi dan hasilnya VALID (UC-01 selesai sukses).',
    'Informasi petugas dan tarif resmi sudah ditampilkan.',
    'Koneksi internet tersedia untuk pembayaran digital.',
])

doc.add_heading('Postconditions (Success)', level=3)
add_bullet_list(doc, [
    'Pembayaran tercatat di database server (tabel payments).',
    'Status pembayaran: "completed".',
    'Struk digital (receipt) ditampilkan dan dapat disimpan.',
    'Nomor transaksi unik di-generate.',
])

doc.add_heading('Postconditions (Failure)', level=3)
add_bullet_list(doc, [
    'Pembayaran tidak diproses, saldo pengguna tidak terpotong.',
    'Pesan error ditampilkan.',
    'Jika gagal setelah potong saldo → refund otomatis (NFR-12).',
])

doc.add_heading('Basic Flow — Pembayaran QRIS', level=3)
add_flow_table(doc,
    ['Langkah', 'Aktor', 'Sistem'],
    [
        ['1', 'Pengguna melihat detail petugas dan tarif. Menekan tombol "Bayar via QRIS".', '—'],
        ['2', '—', 'Sistem menampilkan halaman pembayaran dengan detail: nama petugas, area parkir, tarif, dan durasi.'],
        ['3', 'Pengguna mengonfirmasi jumlah pembayaran dan durasi parkir.', '—'],
        ['4', '—', 'Sistem men-generate kode QRIS untuk pembayaran (REQ-19).'],
        ['5', 'Pengguna memindai kode QRIS menggunakan aplikasi m-banking atau e-wallet.', '—'],
        ['6', '—', 'Sistem memproses transaksi melalui QRIS gateway secara aman (REQ-21).'],
        ['7', '—', 'Sistem menerima konfirmasi pembayaran dari gateway (REQ-22).'],
        ['8', '—', 'Sistem menyimpan transaksi di database dengan status "completed" (REQ-24).'],
        ['9', '—', 'Sistem menampilkan status "Berhasil" (REQ-23) dan struk digital: nomor transaksi, tanggal/waktu, jumlah, metode, nama petugas, area (REQ-25).'],
        ['10', 'Pengguna dapat menyimpan atau menutup struk.', '—'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Alternative Flow: AF-01 — Pembayaran Tunai (Cash)', level=3)
add_alt_table(doc,
    ['Langkah', 'Deskripsi'],
    [
        ['1a', 'Pengguna memilih "Bayar Tunai" atau petugas mencatat pembayaran tunai.'],
        ['2a', 'Pengguna membayar secara langsung kepada petugas.'],
        ['3a', 'Sistem mencatat pembayaran dengan metode "cash" (REQ-20).'],
        ['4a', 'Sistem menampilkan struk digital dengan metode pembayaran "Tunai".'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Alternative Flow: AF-02 — Pembayaran Gagal', level=3)
add_alt_table(doc,
    ['Langkah', 'Deskripsi'],
    [
        ['7b', 'Gateway pembayaran mengembalikan status gagal.'],
        ['8b', 'Sistem menampilkan status "Pembayaran Gagal" dengan pesan error yang jelas.'],
        ['9b', 'Sistem menyimpan transaksi dengan status "failed".'],
        ['10b', 'Pengguna diberi opsi untuk mencoba lagi atau batalkan.'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Alternative Flow: AF-03 — Timeout Pembayaran', level=3)
add_alt_table(doc,
    ['Langkah', 'Deskripsi'],
    [
        ['6c', 'Tidak ada konfirmasi pembayaran dalam waktu yang ditentukan.'],
        ['7c', 'Sistem menampilkan pesan: "Waktu pembayaran habis. Silakan coba lagi."'],
        ['8c', 'Status transaksi: "pending" (dapat di-reconcile nanti).'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Exception Flow', level=3)
add_styled_table(doc,
    ['Kode', 'Deskripsi'],
    [
        ['EX-01', 'Double charge terdeteksi → Sistem menggunakan idempotency key untuk mencegah pembayaran ganda (NFR-12). Refund otomatis jika terjadi.'],
        ['EX-02', 'Koneksi terputus saat transaksi → Sistem menunggu dan melakukan retry. Status transaksi tetap "pending" hingga dikonfirmasi.'],
    ],
    header_color="DC2626"
)
doc.add_paragraph('')

doc.add_heading('Business Rules', level=3)
add_bullet_list(doc, [
    'Semua pembayaran (QRIS dan tunai) wajib tercatat di database (BR-04).',
    'Setiap transaksi memiliki idempotency key (NFR-12).',
    'Rekonsiliasi harian wajib dilakukan (NFR-34, BR-04).',
    'Transaksi pembayaran disimpan minimal 10 tahun (BR-05).',
    'Integrasi QRIS harus sesuai standar Bank Indonesia dan PCI DSS (NFR-34).',
])

doc.add_page_break()

# ----- UC-04 -----
doc.add_heading('UC-04: Dashboard Monitoring Real-time', level=2)

add_attr_table(doc, [
    ('Use Case ID', 'UC-04'),
    ('Nama', 'Dashboard Monitoring Real-time'),
    ('Deskripsi', 'Admin Dishub mengakses dashboard web untuk memonitor aktivitas sistem secara real-time, termasuk statistik scan, data pembayaran, dan laporan masuk. Tampilan publik juga tersedia untuk transparansi.'),
    ('Aktor Utama', 'Admin Dishub (A2)'),
    ('Aktor Pendukung', 'Sistem (A5), Masyarakat Umum (A1 — untuk tampilan publik)'),
    ('Prioritas', 'High'),
    ('Trigger', 'Admin membuka dashboard melalui web browser dan login.'),
    ('Frekuensi', 'Setiap hari — sepanjang jam operasional.'),
])
doc.add_paragraph('')

doc.add_heading('Preconditions', level=3)
add_bullet_list(doc, [
    'Admin telah login ke sistem dengan kredensial valid (role: admin/superadmin).',
    'Browser yang digunakan: Chrome, Firefox, atau Edge versi terbaru.',
    'Resolusi layar minimal 1366x768 piksel.',
    'Koneksi internet stabil tersedia.',
])

doc.add_heading('Postconditions (Success)', level=3)
add_bullet_list(doc, [
    'Dashboard menampilkan data terkini secara real-time.',
    'Admin dapat membuat keputusan berdasarkan data yang ditampilkan.',
])

doc.add_heading('Basic Flow (Main Success Scenario)', level=3)
add_flow_table(doc,
    ['Langkah', 'Aktor', 'Sistem'],
    [
        ['1', 'Admin membuka URL dashboard di web browser.', '—'],
        ['2', 'Admin memasukkan username dan password.', '—'],
        ['3', '—', 'Sistem memvalidasi kredensial dan men-generate JWT token.'],
        ['4', '—', 'Sistem mengambil data dari server (GET /api/dashboard/stats).'],
        ['5', '—', 'Sistem menampilkan: total scan hari ini (REQ-26), jumlah QR aktif vs palsu (REQ-27), grafik tren validasi (REQ-28), heatmap lokasi scan (REQ-29), notifikasi laporan masuk (REQ-30).'],
        ['6', '—', 'Sistem memperbarui data secara real-time (REQ-31).'],
        ['7', 'Admin memfilter data berdasarkan periode waktu (REQ-32).', '—'],
        ['8', 'Admin menganalisis data dan mengambil tindakan.', '—'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Alternative Flow: AF-01 — Tampilan Publik', level=3)
add_alt_table(doc,
    ['Langkah', 'Deskripsi'],
    [
        ['1a', 'Masyarakat umum mengakses URL dashboard publik (REQ-33).'],
        ['2a', 'Sistem menampilkan statistik umum: total scan, jumlah petugas aktif, tren mingguan.'],
        ['3a', 'Data detail (laporan, data personal) TIDAK ditampilkan.'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Alternative Flow: AF-02 — Login Gagal', level=3)
add_alt_table(doc,
    ['Langkah', 'Deskripsi'],
    [
        ['3b', 'Kredensial tidak valid.'],
        ['4b', 'Sistem menampilkan pesan: "Username atau password salah."'],
        ['5b', 'Setelah 5 kali gagal → akun terkunci sementara selama 15 menit.'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Alternative Flow: AF-03 — Session Timeout', level=3)
add_alt_table(doc,
    ['Langkah', 'Deskripsi'],
    [
        ['*', 'Admin tidak aktif selama 30 menit (BR-09, NFR-05).'],
        ['*', 'Sistem secara otomatis logout dan menampilkan halaman login.'],
        ['*', 'Admin harus login ulang untuk melanjutkan.'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Business Rules', level=3)
add_bullet_list(doc, [
    'Dashboard publik dapat diakses tanpa login (BR-06).',
    'Data detail dan data personal tersembunyi di tampilan publik (BR-06).',
    'Session timeout 30 menit untuk admin (BR-09).',
    'Dashboard harus load dalam 3 detik (initial) dan 1 detik (refresh) (NFR-03).',
])

doc.add_page_break()

# ----- UC-05 -----
doc.add_heading('UC-05: Manajemen Akun Petugas & Badge QR', level=2)

add_attr_table(doc, [
    ('Use Case ID', 'UC-05'),
    ('Nama', 'Manajemen Akun Petugas & Badge QR'),
    ('Deskripsi', 'Admin Dishub mendaftarkan petugas parkir baru, men-generate badge QR resmi, dan mengelola status petugas (aktif/non-aktif). Superadmin memiliki akses penuh CRUD.'),
    ('Aktor Utama', 'Admin Dishub (A2), Superadmin (A3)'),
    ('Aktor Pendukung', 'Sistem (A5)'),
    ('Prioritas', 'High'),
    ('Trigger', 'Admin membuka halaman manajemen petugas dan memilih aksi.'),
    ('Frekuensi', 'Berkala — saat ada petugas baru atau perubahan status.'),
])
doc.add_paragraph('')

doc.add_heading('Preconditions', level=3)
add_bullet_list(doc, [
    'Admin/Superadmin telah login ke sistem.',
    'Admin memiliki role minimal "admin" untuk registrasi; role "superadmin" untuk CRUD penuh.',
])

doc.add_heading('Postconditions (Success)', level=3)
add_bullet_list(doc, [
    'Data petugas baru tersimpan di database (tabel officers).',
    'QR code unik ter-generate otomatis dengan format LOHPARKIR-DSH-YYYY-NNN (REQ-35).',
    'Badge QR siap dicetak (REQ-36).',
])

doc.add_heading('Basic Flow — Registrasi Petugas Baru', level=3)
add_flow_table(doc,
    ['Langkah', 'Aktor', 'Sistem'],
    [
        ['1', 'Admin membuka halaman "Manajemen Petugas" di dashboard.', '—'],
        ['2', 'Admin menekan tombol "Tambah Petugas Baru".', '—'],
        ['3', 'Admin mengisi data petugas: nama, NIP, nomor telepon, zona/area penugasan (REQ-34).', '—'],
        ['4', 'Admin menekan tombol "Simpan".', '—'],
        ['5', '—', 'Sistem memvalidasi input (cek duplikasi NIP dan badge number) (REQ-38).'],
        ['6', '—', 'Sistem men-generate QR code unik: LOHPARKIR-DSH-YYYY-NNN (REQ-35).'],
        ['7', '—', 'Sistem menyimpan data petugas ke database (POST /api/officers).'],
        ['8', '—', 'Sistem menampilkan konfirmasi dan preview badge QR yang dapat dicetak (REQ-36).'],
        ['9', 'Admin mencetak badge QR dan menyerahkan kepada petugas.', '—'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Alternative Flow: AF-01 — Nonaktifkan Petugas', level=3)
add_alt_table(doc,
    ['Langkah', 'Deskripsi'],
    [
        ['1a', 'Admin memilih petugas dari daftar.'],
        ['2a', 'Admin menekan tombol "Nonaktifkan" (REQ-37).'],
        ['3a', 'Sistem mengubah status petugas menjadi "inactive".'],
        ['4a', 'QR code petugas tersebut tidak lagi valid saat di-scan (otomatis ditolak).'],
        ['5a', 'Sistem mencatat perubahan di audit trail.'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Alternative Flow: AF-02 — Duplikasi NIP/Badge', level=3)
add_alt_table(doc,
    ['Langkah', 'Deskripsi'],
    [
        ['5c', 'Sistem mendeteksi NIP atau badge number sudah ada di database (REQ-38).'],
        ['6c', 'Sistem menampilkan pesan: "NIP atau Nomor Badge sudah terdaftar."'],
        ['7c', 'Admin memperbaiki data dan mencoba kembali.'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Alternative Flow: AF-03 — Edit Data Petugas (Superadmin)', level=3)
add_alt_table(doc,
    ['Langkah', 'Deskripsi'],
    [
        ['1d', 'Superadmin memilih petugas dan menekan "Edit".'],
        ['2d', 'Superadmin mengubah data yang diperlukan (nama, area, tarif, dll).'],
        ['3d', 'Sistem memvalidasi dan menyimpan perubahan.'],
        ['4d', 'Sistem mencatat perubahan di audit trail.'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Business Rules', level=3)
add_bullet_list(doc, [
    'Hanya admin atau superadmin yang dapat mendaftarkan petugas baru (BR-07).',
    'Petugas tidak dapat mendaftarkan dirinya sendiri (BR-07).',
    'QR code harus unik dan mengikuti format LOHPARKIR-DSH-YYYY-NNN (BR-03).',
    'Setiap perubahan data tercatat di audit trail (NFR-30).',
])

doc.add_page_break()

# ----- UC-06 -----
doc.add_heading('UC-06: Riwayat Pembayaran', level=2)

add_attr_table(doc, [
    ('Use Case ID', 'UC-06'),
    ('Nama', 'Riwayat Pembayaran'),
    ('Deskripsi', 'Masyarakat dapat melihat riwayat pembayaran dan scan parkir yang pernah dilakukan melalui tab Riwayat di aplikasi mobile.'),
    ('Aktor Utama', 'Masyarakat Umum (A1)'),
    ('Aktor Pendukung', 'Sistem (A5)'),
    ('Prioritas', 'Low (future enhancement)'),
    ('Trigger', 'Pengguna membuka tab "Riwayat" di navigasi bawah aplikasi.'),
    ('Frekuensi', 'Sesekali — saat pengguna ingin mengecek histori.'),
])
doc.add_paragraph('')

doc.add_heading('Preconditions', level=3)
add_bullet_list(doc, [
    'Pengguna telah pernah melakukan setidaknya satu pembayaran atau scan.',
    'Aplikasi memiliki akses ke data histori (dari server atau cache lokal).',
])

doc.add_heading('Postconditions (Success)', level=3)
add_bullet_list(doc, [
    'Daftar riwayat pembayaran ditampilkan lengkap dengan detail.',
])

doc.add_heading('Basic Flow (Main Success Scenario)', level=3)
add_flow_table(doc,
    ['Langkah', 'Aktor', 'Sistem'],
    [
        ['1', 'Pengguna membuka tab "Riwayat" dari navigasi bawah.', '—'],
        ['2', '—', 'Sistem mengambil data riwayat pembayaran dari server (GET /api/payments).'],
        ['3', '—', 'Sistem menampilkan daftar pembayaran (REQ-39) dengan detail: tanggal, jumlah, nama petugas, dan status pembayaran (REQ-40).'],
        ['4', 'Pengguna dapat mengetuk item untuk melihat detail struk.', '—'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Alternative Flow: AF-01 — Belum Ada Riwayat', level=3)
add_alt_table(doc,
    ['Langkah', 'Deskripsi'],
    [
        ['3a', 'Tidak ada data pembayaran.'],
        ['4a', 'Sistem menampilkan pesan: "Belum ada riwayat pembayaran."'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Alternative Flow: AF-02 — Mode Offline', level=3)
add_alt_table(doc,
    ['Langkah', 'Deskripsi'],
    [
        ['2b', 'Gagal mengambil data dari server.'],
        ['3b', 'Sistem menampilkan data dari cache lokal (AsyncStorage) dengan label "Data Offline".'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Business Rules', level=3)
add_bullet_list(doc, [
    'Data pembayaran disimpan minimal 10 tahun (BR-05).',
    'Pengguna hanya dapat melihat riwayat miliknya sendiri (berdasarkan device ID).',
])

doc.add_page_break()

# ----- UC-07 -----
doc.add_heading('UC-07: Manajemen Laporan', level=2)

add_attr_table(doc, [
    ('Use Case ID', 'UC-07'),
    ('Nama', 'Manajemen Laporan'),
    ('Deskripsi', 'Admin Dishub mengelola laporan masuk dari masyarakat: melihat daftar laporan, memperbarui status, dan menambahkan catatan tindak lanjut.'),
    ('Aktor Utama', 'Admin Dishub (A2)'),
    ('Aktor Pendukung', 'Sistem (A5), Masyarakat Umum (A1 — sebagai penerima notifikasi)'),
    ('Prioritas', 'High'),
    ('Trigger', 'Admin membuka halaman manajemen laporan di dashboard.'),
    ('Frekuensi', 'Setiap hari — saat ada laporan masuk.'),
])
doc.add_paragraph('')

doc.add_heading('Preconditions', level=3)
add_bullet_list(doc, [
    'Admin telah login ke sistem (role: admin/superadmin).',
    'Terdapat laporan yang telah disubmit oleh masyarakat.',
])

doc.add_heading('Postconditions (Success)', level=3)
add_bullet_list(doc, [
    'Status laporan diperbarui di database.',
    'Notifikasi perubahan status dikirim ke pelapor (UC-08).',
    'Catatan admin tercatat di laporan.',
])

doc.add_heading('Basic Flow (Main Success Scenario)', level=3)
add_flow_table(doc,
    ['Langkah', 'Aktor', 'Sistem'],
    [
        ['1', 'Admin membuka halaman "Manajemen Laporan" di dashboard.', '—'],
        ['2', '—', 'Sistem menampilkan daftar semua laporan dengan opsi filter (REQ-41): berdasarkan status, jenis, dan tanggal.'],
        ['3', 'Admin memilih laporan untuk dilihat detailnya.', '—'],
        ['4', '—', 'Sistem menampilkan detail: foto bukti, lokasi GPS (peta), deskripsi, QR code terkait, tanggal, status saat ini.'],
        ['5', 'Admin memperbarui status laporan (REQ-42): dari "pending" → "in_progress" / "resolved" / "rejected".', '—'],
        ['6', 'Admin menambahkan catatan tindak lanjut (REQ-43).', '—'],
        ['7', 'Admin menekan tombol "Simpan".', '—'],
        ['8', '—', 'Sistem menyimpan perubahan ke database (PUT /api/reports/:id).'],
        ['9', '—', 'Sistem mengirim notifikasi ke pelapor tentang perubahan status (REQ-44). → UC-08.'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Alternative Flow: AF-01 — Cetak Laporan Bulanan', level=3)
add_alt_table(doc,
    ['Langkah', 'Deskripsi'],
    [
        ['*', 'Admin memilih opsi "Cetak Laporan".'],
        ['*', 'Admin memilih periode (bulanan/custom).'],
        ['*', 'Sistem men-generate laporan dalam format yang dapat dicetak.'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Business Rules', level=3)
add_bullet_list(doc, [
    'Laporan tidak dapat dihapus, hanya diubah statusnya (BR-02).',
    'Status yang valid: pending → in_progress → resolved/rejected (BR-02).',
    'Semua perubahan status tercatat di audit trail (NFR-30).',
    'Laporan disimpan minimal 5 tahun (BR-05).',
])

doc.add_page_break()

# ----- UC-08 -----
doc.add_heading('UC-08: Notifikasi Push', level=2)

add_attr_table(doc, [
    ('Use Case ID', 'UC-08'),
    ('Nama', 'Notifikasi Push'),
    ('Deskripsi', 'Sistem mengirimkan notifikasi push kepada pengguna tentang pembaruan status laporan, kepada admin tentang laporan masuk baru, dan alert darurat dari tombol panik petugas.'),
    ('Aktor Utama', 'Sistem (A5)'),
    ('Aktor Pendukung', 'Masyarakat Umum (A1), Admin Dishub (A2), Petugas Parkir (A4)'),
    ('Prioritas', 'Medium'),
    ('Trigger', '(1) Status laporan diubah oleh admin, (2) Laporan baru disubmit, atau (3) Petugas menekan tombol panik darurat.'),
    ('Frekuensi', 'Sering — setiap kali ada perubahan status atau laporan baru.'),
])
doc.add_paragraph('')

doc.add_heading('Preconditions', level=3)
add_bullet_list(doc, [
    'Firebase Cloud Messaging (FCM) telah dikonfigurasi.',
    'Perangkat penerima telah mengizinkan notifikasi push.',
    'Perangkat penerima terhubung ke internet.',
])

doc.add_heading('Postconditions (Success)', level=3)
add_bullet_list(doc, [
    'Notifikasi terkirim ke perangkat tujuan.',
    'Notifikasi tercatat di database (tabel notifications jika tersedia).',
])

doc.add_heading('Skenario 1: Notifikasi Pembaruan Status Laporan', level=3)
add_flow_table(doc,
    ['Langkah', 'Aktor', 'Sistem'],
    [
        ['1', 'Admin memperbarui status laporan (dari UC-07).', '—'],
        ['2', '—', 'Sistem mendeteksi perubahan status.'],
        ['3', '—', 'Sistem mengirim push notification ke perangkat pelapor (REQ-45).'],
        ['4', '—', 'Isi notifikasi: "Laporan [nomor tiket] telah diperbarui. Status: [status baru]."'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Skenario 2: Notifikasi Laporan Baru ke Admin', level=3)
add_flow_table(doc,
    ['Langkah', 'Aktor', 'Sistem'],
    [
        ['1', 'Masyarakat mengirim laporan baru (dari UC-02).', '—'],
        ['2', '—', 'Sistem mendeteksi laporan baru masuk.'],
        ['3', '—', 'Sistem mengirim push notification ke semua admin aktif (REQ-46).'],
        ['4', '—', 'Isi notifikasi: "Laporan baru diterima: [jenis laporan] di [lokasi]."'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Skenario 3: Alert Darurat (Panic Button)', level=3)
add_flow_table(doc,
    ['Langkah', 'Aktor', 'Sistem'],
    [
        ['1', 'Petugas parkir menekan tombol panik darurat di aplikasi officer.', '—'],
        ['2', '—', 'Sistem menangkap lokasi GPS petugas secara real-time (REQ-48).'],
        ['3', '—', 'Sistem mengirim notifikasi PRIORITAS TINGGI ke semua admin (REQ-47).'],
        ['4', '—', 'Isi notifikasi: "DARURAT: Petugas [nama] membutuhkan bantuan di [lokasi GPS]."'],
        ['5', 'Admin menerima notifikasi dan mengambil tindakan segera.', '—'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Alternative Flow: AF-01 — Notifikasi Gagal Terkirim', level=3)
add_alt_table(doc,
    ['Langkah', 'Deskripsi'],
    [
        ['*', 'FCM gagal mengirim notifikasi (perangkat offline, token expired).'],
        ['*', 'Sistem menyimpan notifikasi sebagai "unread" di database.'],
        ['*', 'Notifikasi akan ditampilkan saat pengguna membuka aplikasi (in-app notification).'],
    ]
)
doc.add_paragraph('')

doc.add_heading('Business Rules', level=3)
add_bullet_list(doc, [
    'Alert darurat (panic button) memiliki prioritas tertinggi (BR-10).',
    'Hanya petugas terdaftar yang dapat menggunakan panic button (BR-10).',
    'Notifikasi darurat harus terkirim dalam hitungan detik.',
])

doc.add_page_break()

# ===== 4. MATRIKS TRACEABILITY =====
doc.add_heading('4. Matriks Traceability', level=1)

doc.add_heading('4.1 Use Case → Functional Requirements', level=2)
add_styled_table(doc,
    ['Use Case', 'REQ Coverage'],
    [
        ['UC-01', 'REQ-1, REQ-2, REQ-3, REQ-4, REQ-5, REQ-6, REQ-7, REQ-8, REQ-9'],
        ['UC-02', 'REQ-10, REQ-11, REQ-12, REQ-13, REQ-14, REQ-15, REQ-16, REQ-17, REQ-18'],
        ['UC-03', 'REQ-19, REQ-20, REQ-21, REQ-22, REQ-23, REQ-24, REQ-25'],
        ['UC-04', 'REQ-26, REQ-27, REQ-28, REQ-29, REQ-30, REQ-31, REQ-32, REQ-33'],
        ['UC-05', 'REQ-34, REQ-35, REQ-36, REQ-37, REQ-38'],
        ['UC-06', 'REQ-39, REQ-40'],
        ['UC-07', 'REQ-41, REQ-42, REQ-43, REQ-44'],
        ['UC-08', 'REQ-45, REQ-46, REQ-47, REQ-48'],
    ]
)
doc.add_paragraph('')

doc.add_heading('4.2 Use Case → Non-Functional Requirements', level=2)
add_styled_table(doc,
    ['Use Case', 'NFR Coverage'],
    [
        ['UC-01', 'NFR-01, NFR-02, NFR-07, NFR-08, NFR-17, NFR-24, NFR-26'],
        ['UC-02', 'NFR-01, NFR-08, NFR-10, NFR-11, NFR-13, NFR-17'],
        ['UC-03', 'NFR-01, NFR-06, NFR-12, NFR-14, NFR-34'],
        ['UC-04', 'NFR-03, NFR-05, NFR-06, NFR-31'],
        ['UC-05', 'NFR-05, NFR-07, NFR-09, NFR-30'],
        ['UC-06', 'NFR-01, NFR-17, NFR-33'],
        ['UC-07', 'NFR-05, NFR-09, NFR-30, NFR-33'],
        ['UC-08', 'NFR-15'],
    ]
)
doc.add_paragraph('')

doc.add_heading('4.3 Use Case → Business Rules', level=2)
add_styled_table(doc,
    ['Use Case', 'Business Rules'],
    [
        ['UC-01', 'BR-03'],
        ['UC-02', 'BR-01, BR-02, BR-08'],
        ['UC-03', 'BR-04, BR-05'],
        ['UC-04', 'BR-06, BR-09'],
        ['UC-05', 'BR-03, BR-07'],
        ['UC-06', 'BR-05'],
        ['UC-07', 'BR-02, BR-05'],
        ['UC-08', 'BR-10'],
    ]
)
doc.add_paragraph('')

doc.add_heading('4.4 Aktor → Use Case', level=2)
add_styled_table(doc,
    ['Aktor', 'Use Cases'],
    [
        ['A1 — Masyarakat Umum', 'UC-01, UC-02, UC-03, UC-06'],
        ['A2 — Admin Dishub', 'UC-04, UC-05, UC-07'],
        ['A3 — Superadmin', 'UC-05'],
        ['A4 — Petugas Parkir', 'UC-03, UC-08'],
        ['A5 — Sistem', 'UC-01 s/d UC-08'],
    ]
)

# Footer
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Dokumen ini disusun berdasarkan SRS LohParkir v1.1 (Revised) dan analisis terhadap implementasi sistem yang berjalan.')
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

# Set margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

doc.save('docs/Use_Case_Specification_LohParkir.docx')
print('DOCX generated successfully!')
