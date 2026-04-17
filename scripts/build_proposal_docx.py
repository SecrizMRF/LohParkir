"""Convert proposal markdown into a styled .docx file."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = Path("attached_assets/proposal_lohparkir_bab3-6.md")
OUT = Path("attached_assets/Proposal_LohParkir_BAB3-6.docx")

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def add_heading(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    return h

def add_paragraph_with_inline(text):
    p = doc.add_paragraph()
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r = p.add_run(part[1:-1]); r.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(10)
        else:
            p.add_run(part)
    return p

def add_code_block(lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run("\n".join(lines))
    r.font.name = "Consolas"
    r.font.size = Pt(9)

def add_table(header, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ridx, row in enumerate(rows, 1):
        cells = table.rows[ridx].cells
        for cidx, val in enumerate(row):
            cells[cidx].text = val

lines = SRC.read_text(encoding="utf-8").splitlines()
i = 0
in_code = False
code_buf = []

while i < len(lines):
    line = lines[i]

    if line.startswith("```"):
        if in_code:
            add_code_block(code_buf)
            code_buf = []
            in_code = False
        else:
            in_code = True
        i += 1
        continue
    if in_code:
        code_buf.append(line)
        i += 1
        continue

    if line.startswith("# "):
        add_heading(line[2:].strip(), 0)
    elif line.startswith("## "):
        add_heading(line[3:].strip(), 1)
    elif line.startswith("### "):
        add_heading(line[4:].strip(), 2)
    elif line.startswith("#### "):
        add_heading(line[5:].strip(), 3)
    elif line.startswith("> "):
        p = add_paragraph_with_inline(line[2:].strip())
        p.paragraph_format.left_indent = Cm(0.75)
        for r in p.runs:
            r.italic = True
    elif line.startswith("---"):
        doc.add_paragraph().add_run("―" * 40)
    elif re.match(r"^\d+\.\s", line):
        p = add_paragraph_with_inline(re.sub(r"^\d+\.\s", "", line))
        p.style = doc.styles["List Number"]
    elif line.startswith("- "):
        p = add_paragraph_with_inline(line[2:])
        p.style = doc.styles["List Bullet"]
    elif line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s\-\|:]+\|$", lines[i + 1]):
        header = [c.strip() for c in line.strip("|").split("|")]
        i += 2
        rows = []
        while i < len(lines) and lines[i].startswith("|"):
            rows.append([c.strip() for c in lines[i].strip("|").split("|")])
            i += 1
        add_table(header, rows)
        continue
    elif line.strip() == "":
        pass
    else:
        add_paragraph_with_inline(line)

    i += 1

doc.save(OUT)
print(f"Saved: {OUT}")
